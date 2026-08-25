"""
AMTLICH! Das biometrische Bürgeramt - DOCX Print-to-Play Generator
Erstellt ein vollständiges, formatiertes Word-Dokument (.docx) für den physischen Playtest.
"""

import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Setzt die Hintergrundfarbe einer Tabellenzelle."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Setzt Innenabstände (Padding) einer Zelle."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, color="999999", sz="4", val="dashed"):
    """Setzt Rahmenlinien für Schnittkanten."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def create_docx():
    with open('print-to-play/cards_data.json', 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    # Set page margins to 0.5 inch (compact print)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

    # -------------------------------------------------------------
    # 1. Deckblatt & Materialübersicht
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("🏛️ AMTLICH! Das biometrische Bürgeramt")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Print-to-Play Materialpaket für den physischen Playtest (Spielkonzept v0.2)")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    # Checkliste
    intro = doc.add_paragraph()
    intro.add_run("📦 In diesem Dokument enthaltenes Spielmaterial:\n").bold = True
    items = [
        ("4x Kommunentableaus", "Spieler-Tableaus mit 2 Personalplätzen, 3 Modernisierungsplätzen, Warteschlange und Zeit-Track"),
        ("12x Personalkarten", "Zufällige Schreibtisch-Besetzungen mit individuellen Arbeitsstilen"),
        ("36x Modernisierungskarten", "4 persönliche Sets à 9 Modernisierungen (Digitalisierung, Infrastruktur, Schulung)"),
        ("24x Startdeck-Aktionskarten", "4 identische Startdecks à 6 Karten"),
        ("48x Markt-Aktionskarten", "Gemeinsamer Markt (Hilfe, Organisation, Störfall, Reaktion)"),
        ("48x Antragskarten", "Bürger-Fälle mit biometrischem Flavor-Text (mechanisch einheitlich: 2 Startmarken, 1 Punkt)"),
        ("Token & Markierungsbogen", "Bearbeitungsmarken, Zeitmarker (⏱️), Rundenzähler 1-8 und Startspieler-Stempel"),
        ("Regel-Kurzübersicht", "Kompakter Leitfaden für den Spieltisch")
    ]
    for name, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # -------------------------------------------------------------
    # 2. Regel-Kurzübersicht (Quickplay Guide)
    # -------------------------------------------------------------
    h1 = doc.add_heading("📖 Dienstvorschrift: Kurzspielregel (8 Runden)", level=1)
    
    r_p = doc.add_paragraph()
    r_p.add_run("Ziel des Spiels: ").bold = True
    r_p.add_run("Leite dein Bürgeramt über 8 Runden. Schließe möglichst viele Anträge ab (+1 Punkt) und halte deine Warteschlange unter Kontrolle (Schlussmalus: -1 Punkt je 2 offene Fälle, aufgerundet).\n")

    doc.add_heading("Ablauf eines Spielerzugs (3 Zeitpunkte erhalten):", level=2)
    steps = [
        ("1. Zeit erhalten", "Erhalte 3 Zeit (⏱️). Restzeit verfällt am Zugende."),
        ("2. Anträge annehmen", "Nimm 1 (Pflicht), 2 oder 3 Anträge. Freie Schreibtische werden prioritär mit den ältesten wartenden Fällen besetzt. Neue Fälle starten mit 2 Bearbeitungsmarken. Überschüssige Fälle gehen in die Warteschlange."),
        ("3. Aktionen & Modernisierung", "Spiele Handkarten gegen Zeitkosten aus. Investiere beliebig viel Zeit in bis zu 3 Modernisierungen. Schulungen werden Schreibtisch 1 oder 2 zugeordnet."),
        ("4. Marktkarte erwerben", "Nimm kostenlos bis zu 1 der 3 offenen Marktkarten auf deine persönliche Ablage. Markt wird nachgefüllt.")
    ]
    for s_name, s_desc in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"{s_name}: ").bold = True
        p.add_run(s_desc)

    doc.add_heading("Gemeinsames Rundenende (nachdem alle am Zug waren):", level=2)
    end_steps = [
        "1. Bearbeitung: Von jedem regulär bearbeiteten aktiven Schreibtisch wird 1 Bearbeitungsmarke entfernt.",
        "2. Wertung: Anträge mit 0 Marken wandern in den Wertungsstapel (+1 Punkt).",
        "3. Störfall-Übergabe: Bei Abschluss eines Falls mit fremdem Störfall wählt die betroffene Kommune: Lerneffekt (Karte ins eigene Deck) oder Fall abgeschlossen (entsorgen).",
        "4. Nachbesetzung: Freie Plätze werden aus der Warteschlange nachbesetzt (werden erst am nächsten Rundenende bearbeitet).",
        "5. Handkarten: Alle legen Handkarten ab und ziehen 4 neue Karten.",
        "6. Startspieler: Marker wandert im Uhrzeigersinn weiter."
    ]
    for es in end_steps:
        doc.add_paragraph(es, style='List Bullet')

    doc.add_heading("Warteschlangen-Status & Schwellen:", level=2)
    q_table = doc.add_table(rows=4, cols=3)
    q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Wartende Fälle", "Status", "Allgemeine Wirkung"]
    for i, h in enumerate(headers):
        cell = q_table.cell(0, i)
        cell.text = h
        set_cell_background(cell, "0F172A")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    q_data = [
        ("0–2 Anträge", "Normalbetrieb", "Regulärer Dienstbetrieb."),
        ("3–4 Anträge", "Andrang ⚠️", "Fremde Störfälle gegen dich kosten +1 Zeit. Schaltet Triage-/Ruhe-Fähigkeiten frei."),
        ("5+ Anträge", "Überlastung 🚨", "Wie Andrang; zusätzlich kostet deine erste positive Aktion pro Zug 1 Zeit weniger.")
    ]
    for row_idx, row_vals in enumerate(q_data):
        for col_idx, val in enumerate(row_vals):
            cell = q_table.cell(row_idx + 1, col_idx)
            cell.text = val
            if row_idx % 2 == 1:
                set_cell_background(cell, "F1F5F9")

    doc.add_page_break()

    # -------------------------------------------------------------
    # 3. 4x Kommunentableaus (Player Mats)
    # -------------------------------------------------------------
    for player_num in range(1, 5):
        doc.add_heading(f"🏛️ KOMMUNENTABLEAU: KOMMUNE {player_num}", level=1)
        p = doc.add_paragraph("Lege diese Seite vor dich. Platziere 2 Personalkarten auf die Schreibtische und verwalte hier deine Anträge.\n")

        # Top Track: Zeit & Rundenzähler
        top_table = doc.add_table(rows=2, cols=4)
        top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        top_table.cell(0, 0).text = "⏱️ ZEIT-VORRAT (3 Zeit / Zug)"
        top_table.cell(0, 1).text = "Zeit 1 [ ⏱ ]"
        top_table.cell(0, 2).text = "Zeit 2 [ ⏱ ]"
        top_table.cell(0, 3).text = "Zeit 3 [ ⏱ ]"

        top_table.cell(1, 0).text = "🏆 WERTUNGSSTAPEL"
        top_table.cell(1, 1).text = "(Hier erledigte Anträge ablegen: 1 Pkt. / Karte)"
        top_table.cell(1, 2).text = "🎴 AKTIONISDECK"
        top_table.cell(1, 3).text = "🗄️ ABLAGESTAPEL"
        
        for r in range(2):
            for c in range(4):
                set_cell_background(top_table.cell(r, c), "F8FAFC")
                set_cell_borders(top_table.cell(r, c), "64748B", "6", "solid")

        doc.add_paragraph()

        # Desks Table
        desks_table = doc.add_table(rows=3, cols=2)
        desks_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        desks_table.cell(0, 0).text = "🏢 SCHREIBTISCH 1 (Personalplatz 1)"
        desks_table.cell(0, 1).text = "🏢 SCHREIBTISCH 2 (Personalplatz 2)"
        set_cell_background(desks_table.cell(0, 0), "1E3A8A")
        set_cell_background(desks_table.cell(0, 1), "1E3A8A")
        desks_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        desks_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        desks_table.cell(1, 0).text = "[ Hier Personalkarte 1 platzieren ]\n(Dauerhafte Fähigkeit gilt für Schreibtisch 1)"
        desks_table.cell(1, 1).text = "[ Hier Personalkarte 2 platzieren ]\n(Dauerhafte Fähigkeit gilt für Schreibtisch 2)"

        desks_table.cell(2, 0).text = "📄 AKTIVER ANTRAG (Schreibtisch 1)\nStart: 🔴 🔴 (2 Bearbeitungsmarken)\nStörfall-Platz: (Max. 2 fremde Störfälle anlegbar)"
        desks_table.cell(2, 1).text = "📄 AKTIVER ANTRAG (Schreibtisch 2)\nStart: 🔴 🔴 (2 Bearbeitungsmarken)\nStörfall-Platz: (Max. 2 fremde Störfälle anlegbar)"

        for r in [1, 2]:
            for c in [0, 1]:
                cell = desks_table.cell(r, c)
                set_cell_background(cell, "FFFFFF")
                set_cell_borders(cell, "CBD5E1", "6", "solid")
                set_cell_margins(cell, top=150, bottom=150, left=150, right=150)

        doc.add_paragraph()

        # Warteschlange Area
        q_box = doc.add_table(rows=2, cols=3)
        q_box.alignment = WD_TABLE_ALIGNMENT.CENTER
        q_box.cell(0, 0).text = "⏳ WARTESCHLANGE (Normal: 0-2)"
        q_box.cell(0, 1).text = "⚠️ ANDRANG (3-4 Anträge)"
        q_box.cell(0, 2).text = "🚨 ÜBERLASTUNG (5+ Anträge)"
        set_cell_background(q_box.cell(0, 0), "10B981")
        set_cell_background(q_box.cell(0, 1), "F59E0B")
        set_cell_background(q_box.cell(0, 2), "EF4444")
        for c in range(3):
            q_box.cell(0, c).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            q_box.cell(0, c).paragraphs[0].runs[0].font.bold = True

        q_box.cell(1, 0).text = "Hier wartende Anträge in Reihe anlegen (Älteste zuerst).\nRegulärer Betrieb."
        q_box.cell(1, 1).text = "Gegnerische Störfälle kosten +1 Zeit!\nTriage-Fähigkeiten werden aktiviert."
        q_box.cell(1, 2).text = "Gegner-Störfälle kosten +1 Zeit.\nErste eigene Hilfe-Aktion kostet -1 Zeit!"

        for c in range(3):
            set_cell_borders(q_box.cell(1, c), "CBD5E1", "6", "solid")
            set_cell_margins(q_box.cell(1, c), top=120, bottom=120, left=120, right=120)

        doc.add_paragraph()

        # Modernisierungen
        mod_box = doc.add_table(rows=2, cols=3)
        mod_box.alignment = WD_TABLE_ALIGNMENT.CENTER
        mod_box.cell(0, 0).text = "⚙️ MODERNISIERUNGSPLATZ 1"
        mod_box.cell(0, 1).text = "⚙️ MODERNISIERUNGSPLATZ 2"
        mod_box.cell(0, 2).text = "⚙️ MODERNISIERUNGSPLATZ 3"
        for c in range(3):
            set_cell_background(mod_box.cell(0, c), "334155")
            mod_box.cell(0, c).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            mod_box.cell(0, c).paragraphs[0].runs[0].font.bold = True

        mod_box.cell(1, 0).text = "[ Modernisierung 1 ]\nZeit-Fortschritt hier anlegen.\nSobald bezahlt: DAUERHAFT AKTIV."
        mod_box.cell(1, 1).text = "[ Modernisierung 2 ]\nZeit-Fortschritt hier anlegen.\nSobald bezahlt: DAUERHAFT AKTIV."
        mod_box.cell(1, 2).text = "[ Modernisierung 3 ]\nZeit-Fortschritt hier anlegen.\nSobald bezahlt: DAUERHAFT AKTIV."
        for c in range(3):
            set_cell_borders(mod_box.cell(1, c), "CBD5E1", "6", "solid")
            set_cell_margins(mod_box.cell(1, c), top=120, bottom=120, left=120, right=120)

        doc.add_page_break()

    # -------------------------------------------------------------
    # Helper to render Card Sheets in DOCX
    # -------------------------------------------------------------
    def render_card_grid(card_list, category_title, card_type_default=""):
        doc.add_heading(category_title, level=1)
        p = doc.add_paragraph("Ausschneiden entlang der gestrichelten Linien (Schnittmarken).")
        
        # 3 columns card grid
        rows_needed = (len(card_list) + 2) // 3
        table = doc.add_table(rows=rows_needed, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx, card in enumerate(card_list):
            r = idx // 3
            c = idx % 3
            cell = table.cell(r, c)
            set_cell_borders(cell, "64748B", "4", "dashed")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)

            # Card Header
            title = card.get('name') or card.get('title')
            cost = card.get('cost')
            ctype = card.get('type') or card_type_default
            icon = card.get('icon', '📄')

            cell_p = cell.paragraphs[0]
            header_run = cell_p.add_run(f"[{ctype.upper()}]" + (f"  ⏱️ {cost} Zeit" if cost is not None else "") + "\n")
            header_run.font.size = Pt(8)
            header_run.font.bold = True
            header_run.font.color.rgb = RGBColor(30, 58, 138)

            name_run = cell_p.add_run(f"{icon} {title}\n")
            name_run.font.size = Pt(9.5)
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(15, 23, 42)

            if card.get('tag'):
                tag_run = cell_p.add_run(f"🔖 {card['tag']}\n")
                tag_run.font.size = Pt(7.5)
                tag_run.font.italic = True
                tag_run.font.color.rgb = RGBColor(100, 116, 139)

            if card.get('synergy'):
                tag_run = cell_p.add_run(f"⚡ {card['synergy']}\n")
                tag_run.font.size = Pt(7.5)
                tag_run.font.bold = True
                tag_run.font.color.rgb = RGBColor(217, 119, 6)

            desc = card.get('description') or card.get('flavor') or ""
            desc_run = cell_p.add_run(f"{desc}\n")
            desc_run.font.size = Pt(8)
            if card.get('flavor'):
                desc_run.font.italic = True
                desc_run.font.color.rgb = RGBColor(80, 80, 80)
            else:
                desc_run.font.color.rgb = RGBColor(30, 41, 59)

            if 'baseTokens' in card or 'Start: 2 Bearbeitungsmarken' in desc:
                footer_run = cell_p.add_run("Start: 🔴 🔴 (2 Marken) · Wert: 🏆 1 Punkt")
                footer_run.font.size = Pt(7.5)
                footer_run.font.bold = True

        doc.add_page_break()

    # 4. Personalkarten (12)
    render_card_grid(data['staff'], "👨‍💼 PERSONALKARTEN (12 Profile)", "Personal")

    # 5. Modernisierungskarten (4 Sets à 9)
    for set_idx in range(1, 5):
        render_card_grid(data['modernizations'], f"⚙️ MODERNISIERUNGSVORRAT – SET {set_idx} (9 Karten)", "Modernisierung")

    # 6. Startdeckkarten (4 Sets à 6)
    for set_idx in range(1, 5):
        render_card_grid(data['startDeck'], f"🎴 STARTDECK – SPIELER {set_idx} (6 Karten)", "Aktion")

    # 7. Marktkarten (48 Karten)
    all_market = []
    for c in data['marketCards']:
        copies = c.get('copies', 1)
        for i in range(copies):
            all_market.append(c)
    render_card_grid(all_market, "🛒 ZENTRALER AKTIONSMARKT (48 Karten)", "Markt")

    # 8. Antragskarten (48 Karten)
    render_card_grid(data['requests'], "📄 ANTRAGSKARTEN (48 Bürger-Fälle)", "Antrag")

    # 9. Token & Marker Sheet
    doc.add_heading("🪙 TOKEN- & MARKIERUNGSBOGEN (Ausschneiden)", level=1)
    doc.add_paragraph("Kleben Sie dieses Blatt auf Pappe und schneiden Sie die Marker aus:")

    token_table = doc.add_table(rows=6, cols=6)
    token_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tokens_list = []
    # 24 Bearbeitungsmarken
    for i in range(18):
        tokens_list.append(("🔴 MARKE", "1 Bearbeitungsmarke", "FEE2E2"))
    # 12 Zeitmarker
    for p in range(1, 5):
        for t in range(1, 4):
            tokens_list.append((f"⏱️ ZEIT {t}", f"Kommune {p}", "DBEAFE"))
    # Rundenzähler 1-8
    for r in range(1, 9):
        tokens_list.append((f"📅 RUNDE {r}", "Rundenzähler", "FEF3C7"))
    # Startspieler & Spezial
    tokens_list.append(("👑 STARTSPIELER", "Amtlicher Stempel", "FDE68A"))
    tokens_list.append(("🛡️ SCHUTZSCHILD", "Plausibilität", "D1FAE5"))

    for idx, (t_title, t_sub, color) in enumerate(tokens_list[:36]):
        r = idx // 6
        c = idx % 6
        cell = token_table.cell(r, c)
        set_cell_background(cell, color)
        set_cell_borders(cell, "64748B", "4", "solid")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(f"{t_title}\n")
        run1.font.size = Pt(8.5)
        run1.font.bold = True
        run2 = p.add_run(f"{t_sub}")
        run2.font.size = Pt(7)

    output_path = 'print-to-play/AMTLICH_Print_to_Play_v0.2.docx'
    doc.save(output_path)
    print(f"DOCX erfolgreich erstellt: {output_path}")

if __name__ == '__main__':
    create_docx()
